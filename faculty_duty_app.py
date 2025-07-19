from fastapi import FastAPI, UploadFile, File, Request
from fastapi.middleware.cors import CORSMiddleware
import os
import pandas as pd
import json
from fastapi.responses import StreamingResponse, JSONResponse, FileResponse
import tempfile
from io import BytesIO
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.table import WD_ROW_HEIGHT_RULE
import datetime

app = FastAPI()

# Enable CORS for all origins (for development)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

FAKE_FACULTY_PATH = "faculty_upload.csv"
FACULTY_GROUPS_PATH = "faculty_groups.json"
FACULTY_UNAVAILABILITY_PATH = "faculty_unavailability.json"
EXAM_SCHEDULE_PATH = "exam_schedule.json"
ASSIGNMENTS_PATH = "assignments.json"
EXAM_CONFIG_PATH = "exam_config.json"

# Advanced Report Generation Functions
def generate_faculty_summary_excel(df, unavailability=None):
    try:
        print(f"Starting Excel generation with df shape: {df.shape}")
        print(f"DataFrame columns: {df.columns.tolist()}")
        print(f"DataFrame head: {df.head()}")
        
        faculty_list = None
        faculty_contacts = {}
        
        if os.path.exists(FAKE_FACULTY_PATH):
            faculty_df = pd.read_csv(FAKE_FACULTY_PATH)
            faculty_list = faculty_df['faculty'].tolist() if 'faculty' in faculty_df.columns else faculty_df['Faculty'].tolist()
            print(f"Loaded faculty list: {len(faculty_list)} faculty")
            
            # Create a mapping of faculty name to contact info
            for _, row in faculty_df.iterrows():
                faculty_name = row.get('faculty') or row.get('Faculty')
                if faculty_name:
                    faculty_contacts[faculty_name] = {
                        'Phone No': row.get('Phone No', ''),
                        'Email Id': row.get('Email Id', '') or row.get('Email ID', '') or row.get('email', '')
                    }
            print(f"Loaded contact info for {len(faculty_contacts)} faculty")
        
        # Normalize column names to handle both lowercase and uppercase
        column_mapping = {}
        for col in df.columns:
            if col.lower() == 'date':
                column_mapping[col] = 'Date'
            elif col.lower() == 'shift':
                column_mapping[col] = 'Shift'
            elif col.lower() == 'faculty':
                column_mapping[col] = 'Faculty'
        
        # Rename columns if needed
        if column_mapping:
            df = df.rename(columns=column_mapping)
            print(f"Renamed columns: {column_mapping}")
        
        print(f"After column normalization - DataFrame columns: {df.columns.tolist()}")
        
        required_columns = {'Faculty', 'Date', 'Shift'}
        if df is None or df.empty or (set(df.columns) & required_columns) != required_columns:
            print(f"Missing required columns. Required: {required_columns}, Available: {set(df.columns)}")
            df = pd.DataFrame({col: pd.Series(dtype='object') for col in ['Faculty', 'Date', 'Shift']})
        else:
            df = df.copy()
            df['Date'] = pd.to_datetime(df['Date'], errors='coerce').dt.date
            print(f"Processed dates. DataFrame shape: {df.shape}")
        
        if faculty_list is not None:
            print(f"Processing {len(faculty_list)} faculty from list")
            summary_rows = []
            for faculty in faculty_list:
                faculty_df = df[df['Faculty'] == faculty]
                first_half_count = (faculty_df['Shift'] == 'First Half').sum()
                second_half_count = (faculty_df['Shift'] == 'Second Half').sum()
                total_duties = len(faculty_df)
                fh_list = faculty_df.loc[faculty_df['Shift'] == 'First Half', 'Date'].to_list()
                sh_list = faculty_df.loc[faculty_df['Shift'] == 'Second Half', 'Date'].to_list()
                fh_dates = pd.Series(list(pd.to_datetime(fh_list, errors='coerce')), dtype='datetime64[ns]')
                sh_dates = pd.Series(list(pd.to_datetime(sh_list, errors='coerce')), dtype='datetime64[ns]')
                
                if not fh_dates.empty and pd.api.types.is_datetime64_any_dtype(fh_dates):
                    first_half_dates = ', '.join(fh_dates.dt.strftime('%d-%m-%Y'))
                else:
                    first_half_dates = ''
                
                if not sh_dates.empty and pd.api.types.is_datetime64_any_dtype(sh_dates):
                    second_half_dates = ', '.join(sh_dates.dt.strftime('%d-%m-%Y'))
                else:
                    second_half_dates = ''
                
                if unavailability:
                    fh_unavail = format_unavail_dates(unavailability.get(faculty, {'first_half': set()})['first_half'])
                    sh_unavail = format_unavail_dates(unavailability.get(faculty, {'second_half': set()})['second_half'])
                    total_unavail = len(unavailability.get(faculty, {'first_half': set(), 'second_half': set()})['first_half']) + \
                                   len(unavailability.get(faculty, {'first_half': set(), 'second_half': set()})['second_half'])
                else:
                    fh_unavail = 'None'
                    sh_unavail = 'None'
                    total_unavail = 0
                
                # Get contact information for this faculty
                contact_info = faculty_contacts.get(faculty, {})
                phone_no = contact_info.get('Phone No', '')
                email_id = contact_info.get('Email Id', '')
                
                summary_rows.append({
                    'Faculty': faculty,
                    'Phone No': phone_no,
                    'Email ID': email_id,
                    'First Half Duties': first_half_count,
                    'Second Half Duties': second_half_count,
                    'Total Duties': total_duties,
                    'First Half Dates': first_half_dates,
                    'Second Half Dates': second_half_dates,
                    'First Half Unavailable': fh_unavail,
                    'Second Half Unavailable': sh_unavail,
                    'Total Unavailable Slots': total_unavail
                })
            faculty_summary = pd.DataFrame(summary_rows)
            print(f"Created summary DataFrame with {len(faculty_summary)} rows")
        else:
            print("No faculty list found, using groupby approach")
            faculty_summary = df.groupby('Faculty').agg({
                'Shift': lambda x: [
                    sum(x == 'First Half'),
                    sum(x == 'Second Half'),
                    len(x)
                ]
            }).reset_index()
            faculty_summary[['First Half Duties', 'Second Half Duties', 'Total Duties']] = pd.DataFrame(
                faculty_summary['Shift'].tolist(), 
                index=faculty_summary.index
            )
            faculty_summary.drop('Shift', axis=1, inplace=True)
            
            # Add contact information columns
            faculty_summary['Phone No'] = faculty_summary['Faculty'].apply(lambda x: faculty_contacts.get(x, {}).get('Phone No', ''))
            faculty_summary['Email ID'] = faculty_summary['Faculty'].apply(lambda x: faculty_contacts.get(x, {}).get('Email Id', ''))
            
            def get_shift_dates(faculty, shift):
                dates = df[(df['Faculty'] == faculty) & (df['Shift'] == shift)]['Date']
                if not isinstance(dates, pd.Series):
                    dates = pd.Series(dates)
                if not pd.api.types.is_datetime64_any_dtype(dates):
                    dates = pd.to_datetime(dates, errors='coerce')
                dates = dates.dropna()
                if not dates.empty:
                    return ', '.join(dates.dt.strftime('%d-%m-%Y'))
                else:
                    return ''
            
            faculty_summary['First Half Dates'] = faculty_summary['Faculty'].apply(
                lambda f: get_shift_dates(f, 'First Half')
            )
            faculty_summary['Second Half Dates'] = faculty_summary['Faculty'].apply(
                lambda f: get_shift_dates(f, 'Second Half')
            )
            
            if unavailability:
                faculty_summary['First Half Unavailable'] = faculty_summary['Faculty'].apply(
                    lambda f: format_unavail_dates(unavailability.get(f, {'first_half': set()})['first_half'])
                )
                faculty_summary['Second Half Unavailable'] = faculty_summary['Faculty'].apply(
                    lambda f: format_unavail_dates(unavailability.get(f, {'second_half': set()})['second_half'])
                )
                faculty_summary['Total Unavailable Slots'] = faculty_summary['Faculty'].apply(
                    lambda f: len(unavailability.get(f, {'first_half': set(), 'second_half': set()})['first_half']) +
                            len(unavailability.get(f, {'first_half': set(), 'second_half': set()})['second_half'])
                )
            else:
                faculty_summary['First Half Unavailable'] = 'None'
                faculty_summary['Second Half Unavailable'] = 'None'
                faculty_summary['Total Unavailable Slots'] = 0
        
        print(f"Final faculty_summary shape: {faculty_summary.shape}")
        print(f"Final faculty_summary columns: {faculty_summary.columns.tolist()}")
        
        output = BytesIO()
        with pd.ExcelWriter(output, engine="openpyxl") as writer:
            faculty_summary.to_excel(writer, sheet_name="Faculty Duty Summary", index=False)
            worksheet = writer.sheets["Faculty Duty Summary"]
            for idx, col in enumerate(faculty_summary.columns):
                max_length = max(
                    faculty_summary[col].astype(str).apply(len).max(),
                    len(col)
                ) + 2
                worksheet.column_dimensions[chr(65 + idx)].width = min(max_length, 50)
        output.seek(0)
        print("Excel file generated successfully")
        return output.getvalue()
    except Exception as e:
        print(f"Error generating Excel: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

def format_unavail_dates(date_set):
    formatted = []
    for d in sorted(date_set):
        if isinstance(d, str):
            try:
                d_obj = pd.to_datetime(d, errors='coerce')
                if pd.isnull(d_obj):
                    formatted.append(str(d))
                else:
                    formatted.append(d_obj.strftime('%d-%m-%Y'))
            except Exception:
                formatted.append(str(d))
        elif hasattr(d, 'strftime'):
            formatted.append(d.strftime('%d-%m-%Y'))
        else:
            formatted.append(str(d))
    return ', '.join(formatted) if formatted else 'None'

def generate_word_doc(df):
    try:
        df = df.copy()
        
        # Load exam configuration
        exam_config = {
            "examType": "MID SEM",
            "semester": "MO",
            "year": "2025",
            "department": "Computer Science & Engineering",
            "institute": "BIT MESRA, RANCHI"
        }
        if os.path.exists(EXAM_CONFIG_PATH):
            try:
                with open(EXAM_CONFIG_PATH, "r", encoding="utf-8") as f:
                    exam_config = json.load(f)
            except Exception as e:
                print(f"Error loading exam config: {e}")
        
        # Load faculty contact information if available
        faculty_contacts = {}
        if os.path.exists(FAKE_FACULTY_PATH):
            try:
                faculty_df = pd.read_csv(FAKE_FACULTY_PATH)
                # Create a mapping of faculty name to contact info
                for _, row in faculty_df.iterrows():
                    faculty_name = row.get('faculty') or row.get('Faculty')
                    if faculty_name:
                        faculty_contacts[faculty_name] = {
                            'Phone No': row.get('Phone No', ''),
                            'Email Id': row.get('Email Id', '') or row.get('Email ID', '') or row.get('email', '')
                        }
                print(f"Loaded contact info for {len(faculty_contacts)} faculty")
            except Exception as e:
                print(f"Error loading faculty contacts: {e}")
        
        # Normalize column names to handle both lowercase and uppercase
        column_mapping = {}
        for col in df.columns:
            if col.lower() == 'date':
                column_mapping[col] = 'Date'
            elif col.lower() == 'shift':
                column_mapping[col] = 'Shift'
            elif col.lower() == 'faculty':
                column_mapping[col] = 'Faculty'
            elif col.lower() == 'phone no':
                column_mapping[col] = 'Phone No'
            elif col.lower() == 'email id':
                column_mapping[col] = 'Email Id'
        
        # Rename columns if needed
        if column_mapping:
            df = df.rename(columns=column_mapping)
        
        # Add contact information to the dataframe
        df['Phone No'] = df['Faculty'].apply(lambda x: faculty_contacts.get(x, {}).get('Phone No', ''))
        df['Email Id'] = df['Faculty'].apply(lambda x: faculty_contacts.get(x, {}).get('Email Id', ''))
        
        # Always keep Date as datetime.date for logic, only format for display
        df["Date"] = pd.to_datetime(df["Date"], errors="coerce").dt.date
        # Drop rows where Date is missing
        df = df[df["Date"].notna()]
        
        if df.empty:
            print("Warning: No valid data for Word report generation")
            return None
            
        doc = Document()
        # Set narrow margins (0.5 inches = 12.7 mm)
        sections = doc.sections
        for section in sections:
            section.top_margin = Mm(12.7)
            section.bottom_margin = Mm(12.7)
            section.left_margin = Mm(12.7)
            section.right_margin = Mm(12.7)
        
        # Get exam type, semester and year from config
        exam_type = exam_config.get("examType", "MID SEM")
        semester = exam_config.get("semester", "MO")
        year = exam_config.get("year", "2025")
        department = exam_config.get("department", "Computer Science & Engineering")
        institute = exam_config.get("institute", "BIT MESRA, RANCHI")
        
        # Add current date at top right
        today_str = datetime.datetime.today().strftime('%d/%m/%Y')
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(f'Date : {today_str}')
        date_para.alignment = 2  # Right align
        date_run.font.name = 'Times New Roman'
        date_run.font.size = Pt(12)
        
        # Create header with exam type, semester and year
        p = doc.add_paragraph()
        run1 = p.add_run(f"Department of {department}")
        run1.font.name = 'Times New Roman'
        run1.font.size = Pt(12)
        run1.add_break()
        run2 = p.add_run(institute)
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(12)
        p.alignment = 1  # Center align (optional)
        
        header = f"Examination Duty Chart - {exam_type} {semester} {year}"
        p2 = doc.add_paragraph()
        run_header = p2.add_run(header)
        run_header.font.name = 'Times New Roman'
        run_header.font.size = Pt(14)
        run_header.bold = True  # Make header bold
        p2.alignment = 1  # Center align (optional)
        # Add underline to header
        run_header.font.underline = True
        
        # Add time paragraph based on exam_type
        if exam_type == "MID SEM":
            time_text = (
                "Time: 09.40 A.M. to 12.00 NOON (1st Half)\n"
                "01.40 P.M. to 04.00 P.M. (2nd Half)"
            )
        else:
            time_text = (
                "Time: 09.40 A.M. to 01.00 P.M. (1st Half)\n"
                "01.40 P.M. to 05.00 P.M. (2nd Half)"
            )
        p3 = doc.add_paragraph()
        run_time = p3.add_run(time_text)
        run_time.font.name = 'Times New Roman'
        run_time.font.size = Pt(12)
        p3.alignment = 1  # Center align (optional)
        
        # Get unique dates and sort them
        unique_dates = sorted(df["Date"].unique())
        for date in unique_dates:
            if pd.isna(date):
                continue
            # Add date display before the table as bold paragraph
            doc.add_paragraph()  # Add some space before the date display
            date_para = doc.add_paragraph()
            day_str = date.strftime('%A')
            date_str = date.strftime('%d.%m.%Y')
            date_run = date_para.add_run(f"{date_str} ({day_str})")
            date_run.bold = True
            date_run.font.name = 'Times New Roman'
            date_run.font.size = Pt(12)
            
            # Filter DataFrame for the current date
            df_for_date = df[df["Date"] == date].copy()
            # Sort by shift for correct merging order
            df_for_date = df_for_date.sort_values(by=["Shift"])
            
            # Calculate total rows needed including blank row between shifts
            total_rows = len(df_for_date) + 1  # +1 for header
            if "First Half" in df_for_date["Shift"].values and "Second Half" in df_for_date["Shift"].values:
                total_rows += 1  # Add one more row for blank row between shifts
            
            # Create table with calculated rows
            table = doc.add_table(rows=total_rows, cols=5)
            table.style = "Table Grid"
            
            # Set column widths (in mm)
            table.columns[0].width = Mm(25)  # Shift column
            table.columns[1].width = Mm(15.1)  # S.No column (1.51 cm = 15.1 mm)
            table.columns[2].width = Mm(60)  # Faculty column
            table.columns[4].width = Mm(40)  # Email ID column
            table.columns[3].width = Mm(30)  # Phone No column
            table.columns[4].width = Mm(40)  # Email ID column
            
            # Add header row
            hdr_cells = table.rows[0].cells
            headers = ["Shift", "S.No", "Faculty", "Phone No", "Email ID"]
            for i, header in enumerate(headers):
                cell = hdr_cells[i]
                cell.text = header
                # Set all header cells center aligned
                cell.paragraphs[0].alignment = 1  # Center align
                # Make text bold
                for run in cell.paragraphs[0].runs:
                    run.bold = True
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
            
            # Add data rows for this date
            current_row = 1  # Start from row 1 (after header)
            serial_no = 1  # Reset serial number for each date's table
            first_half_rows = []
            second_half_rows = []
            
            # Collect rows for each shift
            for idx, row in df_for_date.iterrows():
                if row["Shift"] == "First Half":
                    first_half_rows.append(row)
                elif row["Shift"] == "Second Half":
                    second_half_rows.append(row)
            
            # Write First Half rows
            for i, row in enumerate(first_half_rows):
                while current_row >= len(table.rows):
                    table.add_row()
                row_cells = table.rows[current_row].cells
                row_cells[0].text = str(row["Shift"]) if i == 0 else ""
                row_cells[1].text = str(serial_no)
                row_cells[2].text = str(row["Faculty"])
                row_cells[3].text = str(row.get("Phone No", ""))
                row_cells[4].text = str(row.get("Email Id", ""))
                for j, cell in enumerate(row_cells):
                    if j in [2, 4]:
                        cell.paragraphs[0].alignment = 0  # Left align
                    else:
                        cell.paragraphs[0].alignment = 1  # Center align
                    for run in cell.paragraphs[0].runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)
                serial_no += 1
                current_row += 1
            first_half_end_row = current_row - 1 if first_half_rows else None
            
            # Add a single blank row if both shifts exist
            if first_half_rows and second_half_rows:
                while current_row >= len(table.rows):
                    table.add_row()
                blank_cells = table.rows[current_row].cells
                for cell in blank_cells:
                    cell.text = ""
                current_row += 1
            
            # Write Second Half rows
            serial_no = 1
            second_half_start_row = current_row
            for i, row in enumerate(second_half_rows):
                while current_row >= len(table.rows):
                    table.add_row()
                row_cells = table.rows[current_row].cells
                row_cells[0].text = str(row["Shift"]) if i == 0 else ""
                row_cells[1].text = str(serial_no)
                row_cells[2].text = str(row["Faculty"])
                row_cells[3].text = str(row.get("Phone No", ""))
                row_cells[4].text = str(row.get("Email Id", ""))
                for j, cell in enumerate(row_cells):
                    if j in [2, 4]:
                        cell.paragraphs[0].alignment = 0  # Left align
                    else:
                        cell.paragraphs[0].alignment = 1  # Center align
                    for run in cell.paragraphs[0].runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)
                serial_no += 1
                current_row += 1
            
            # Merge cells for First Half
            if first_half_rows:
                try:
                    merged_cell = table.cell(1, 0).merge(table.cell(first_half_end_row, 0))
                    merged_cell.vertical_alignment = WD_ROW_HEIGHT_RULE.AT_LEAST
                except Exception as e:
                    print(f"Error merging First Half cells: {e}")
            
            # Merge cells for Second Half
            if second_half_rows:
                try:
                    merged_cell = table.cell(second_half_start_row, 0).merge(table.cell(current_row - 1, 0))
                    merged_cell.vertical_alignment = WD_ROW_HEIGHT_RULE.AT_LEAST
                except Exception as e:
                    print(f"Error merging Second Half cells: {e}")
        
        # Add a note section at the end
        doc.add_paragraph()
        doc.add_heading("Note:", level=1)
        notes = [
            "All the Invigilators according to the invigilation chart are requested to report to the upstairs examination office 20 minute before the examination starts (The room allotment will be done before the start of each examination).",
            "If any Invigilator is unable to do invigilation duty for any reason, then it should be brought to the notice of the Controller of Examination with alternative arrangement through HoD well before the start of the examination.",
            "Invigilators will be prohibited from carrying and using cell phones in the Examination Hall (As recommended in the 66th meeting of the Examination Committee meeting).",
            "Invigilators should make sure that bags of the students are not kept inside the Examination Hall (As recommended in the 66th meeting of the Examination Committee meeting)."
        ]
        for i, note in enumerate(notes, 1):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(0)
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.tab_stops.add_tab_stop(Pt(36))  # Hanging indent at 0.5 inch
            run = p.add_run(f"{i}.\t")
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run2 = p.add_run(note)
            run2.font.name = 'Times New Roman'
            run2.font.size = Pt(12)
            # Hanging indent effect
            p.paragraph_format.first_line_indent = -Pt(18)
            p.paragraph_format.left_indent = Pt(36)
        
        # Add signature section
        doc.add_paragraph("\n\n")
        signature = doc.add_paragraph()
        signature.add_run("(Dr. A. Mustafi)\n").bold = True
        signature.add_run("Professor & Head\n")
        signature.add_run("Department of Computer Science & Engineering\n")
        signature.add_run("B.I.T., Mesra, Ranchi")
        
        # Add copy to section
        doc.add_paragraph("\n")
        copy_to = [
            "All faculty members (through email)",
            "Controller of examination", 
            "Copy to V.C Office",
            "Office File"
        ]
        p_heading = doc.add_paragraph()
        p_heading.add_run("Copy to:").bold = True
        p_recipients = doc.add_paragraph()
        for i, recipient in enumerate(copy_to, 1):
            run = p_recipients.add_run(f"{i}.\t{recipient}")
            if i < len(copy_to):
                run.add_break() # Add a line break instead of a new paragraph
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    except Exception as e:
        print(f"Error generating Word document: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

@app.get("/ping")
def ping():
    return {"message": "pong"}

@app.get("/faculty")
def get_faculty():
    if os.path.exists(FAKE_FACULTY_PATH):
        df = pd.read_csv(FAKE_FACULTY_PATH)
        return df.to_dict(orient="records")
    return []

@app.post("/upload-faculty")
def upload_faculty(file: UploadFile = File(...)):
    filename = file.filename
    temp_path = "temp_upload"
    with open(temp_path, "wb") as f:
        f.write(file.file.read())
    try:
        if filename.endswith('.xlsx'):
            df = pd.read_excel(temp_path)
        elif filename.endswith('.csv'):
            df = pd.read_csv(temp_path, encoding='utf-8', errors='replace')
        else:
            os.remove(temp_path)
            return {"status": "error", "message": "Unsupported file type"}
        df.to_csv("faculty_upload.csv", index=False)
    except Exception as e:
        os.remove(temp_path)
        return {"status": "error", "message": str(e)}
    os.remove(temp_path)
    return {"status": "ok"}

@app.get("/exam-schedule")
def get_exam_schedule():
    if os.path.exists(EXAM_SCHEDULE_PATH):
        with open(EXAM_SCHEDULE_PATH, "r", encoding="utf-8") as f:
            return json.load(f)
    return []

@app.post("/exam-schedule")
def add_exam_schedule(item: dict):
    schedule = []
    if os.path.exists(EXAM_SCHEDULE_PATH):
        with open(EXAM_SCHEDULE_PATH, "r", encoding="utf-8") as f:
            schedule = json.load(f)
    schedule.append(item)
    with open(EXAM_SCHEDULE_PATH, "w", encoding="utf-8") as f:
        json.dump(schedule, f, ensure_ascii=False)
    return {"status": "ok"}

@app.delete("/exam-schedule/{date}")
def delete_exam_schedule(date: str):
    if not os.path.exists(EXAM_SCHEDULE_PATH):
        return {"status": "ok"}
    with open(EXAM_SCHEDULE_PATH, "r", encoding="utf-8") as f:
        schedule = json.load(f)
    # Remove by date (string match)
    schedule = [item for item in schedule if str(item.get('date')) != date]
    with open(EXAM_SCHEDULE_PATH, "w", encoding="utf-8") as f:
        json.dump(schedule, f, ensure_ascii=False)
    return {"status": "ok"}

@app.post("/generate-assignments")
def generate_assignments(data: dict):
    # For demo: just assign first faculty to each slot
    faculty = data.get("faculty", [])
    schedule = data.get("schedule", [])
    assignments = []
    faculty_names = [f.get("faculty") or f.get("Faculty") for f in faculty]
    idx = 0
    for day in schedule:
        for shift, label in [("First Half", "first_half"), ("Second Half", "second_half")]:
            required = day.get(label, 0)
            for _ in range(required):
                if faculty_names:
                    assignments.append({
                        "date": day["date"],
                        "shift": shift,
                        "faculty": faculty_names[idx % len(faculty_names)]
                    })
                    idx += 1
    # Save assignments for report generation
    with open(ASSIGNMENTS_PATH, "w", encoding="utf-8") as f:
        json.dump(assignments, f, ensure_ascii=False)
    return assignments

@app.get("/download-report")
def download_report(type: str):
    try:
        # Load assignments
        if not os.path.exists(ASSIGNMENTS_PATH):
            return JSONResponse(status_code=404, content={"error": "No assignments found"})
        with open(ASSIGNMENTS_PATH, "r", encoding="utf-8") as f:
            assignments = json.load(f)
        df = pd.DataFrame(assignments)
        
        print(f"Download report requested for type: {type}")
        print(f"Assignments data shape: {df.shape}")
        print(f"Assignments columns: {df.columns.tolist()}")
        
        # Load faculty unavailability for Excel report
        unavailability = {}
        if os.path.exists(FACULTY_UNAVAILABILITY_PATH):
            with open(FACULTY_UNAVAILABILITY_PATH, "r", encoding="utf-8") as f:
                unavailability = json.load(f)

        if type == "excel":
            print("Generating Excel report...")
            excel_data = generate_faculty_summary_excel(df, unavailability)
            if excel_data:
                print("Excel report generated successfully")
                return StreamingResponse(
                    BytesIO(excel_data),
                    media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    headers={"Content-Disposition": "attachment; filename=faculty_summary.xlsx"}
                )
            else:
                print("Failed to generate Excel report")
                return JSONResponse(status_code=500, content={"error": "Failed to generate Excel report"})
        elif type == "word":
            print("Generating Word report...")
            word_data = generate_word_doc(df)
            if word_data:
                print("Word report generated successfully")
                return StreamingResponse(
                    BytesIO(word_data),
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={"Content-Disposition": "attachment; filename=faculty_duty_assignment.docx"}
                )
            else:
                print("Failed to generate Word report")
                return JSONResponse(status_code=500, content={"error": "Failed to generate Word report"})
        else:
            return JSONResponse(status_code=400, content={"error": "Invalid report type"})
    except Exception as e:
        print(f"Error in download_report: {str(e)}")
        import traceback
        traceback.print_exc()
        return JSONResponse(status_code=500, content={"error": f"Internal server error: {str(e)}"})

@app.get("/assignments")
def get_assignments():
    if os.path.exists(ASSIGNMENTS_PATH):
        with open(ASSIGNMENTS_PATH, "r", encoding="utf-8") as f:
            return json.load(f)
    return []

@app.post("/assignments")
async def save_assignments(request: Request):
    assignments = await request.json()
    with open(ASSIGNMENTS_PATH, "w", encoding="utf-8") as f:
        json.dump(assignments, f, ensure_ascii=False)
    return {"status": "ok"}

@app.get("/faculty-groups")
def get_faculty_groups():
    if os.path.exists(FACULTY_GROUPS_PATH):
        with open(FACULTY_GROUPS_PATH, "r", encoding="utf-8") as f:
            return json.load(f)
    return []

@app.post("/faculty-groups")
async def save_faculty_groups(request: Request):
    groups = await request.json()
    with open(FACULTY_GROUPS_PATH, "w", encoding="utf-8") as f:
        json.dump(groups, f, ensure_ascii=False)
    return {"status": "ok"}

@app.get("/faculty-unavailability")
def get_faculty_unavailability():
    if os.path.exists(FACULTY_UNAVAILABILITY_PATH):
        with open(FACULTY_UNAVAILABILITY_PATH, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}

@app.post("/faculty-unavailability")
async def save_faculty_unavailability(request: Request):
    unavailability = await request.json()
    with open(FACULTY_UNAVAILABILITY_PATH, "w", encoding="utf-8") as f:
        json.dump(unavailability, f, ensure_ascii=False)
    return {"status": "ok"}

@app.get("/exam-config")
def get_exam_config():
    if os.path.exists(EXAM_CONFIG_PATH):
        with open(EXAM_CONFIG_PATH, "r", encoding="utf-8") as f:
            return json.load(f)
    return {
        "examType": "MID SEM",
        "semester": "MO", 
        "year": "2025",
        "department": "Computer Science & Engineering",
        "institute": "BIT MESRA, RANCHI"
    }

@app.post("/exam-config")
async def save_exam_config(request: Request):
    config = await request.json()
    with open(EXAM_CONFIG_PATH, "w", encoding="utf-8") as f:
        json.dump(config, f, ensure_ascii=False)
    return {"status": "ok"}

@app.post("/regenerate-from-summary")
async def regenerate_from_summary(summary_file: UploadFile = File(...), schedule_file: UploadFile = File(None), unavailability_file: UploadFile = File(None)):
    try:
        print(f"Starting regeneration process...")
        print(f"Summary file: {summary_file.filename if summary_file else 'None'}")
        print(f"Schedule file: {schedule_file.filename if schedule_file else 'None'}")
        print(f"Unavailability file: {unavailability_file.filename if unavailability_file else 'None'}")
        
        if not summary_file.filename.endswith('.xlsx'):
            return JSONResponse(status_code=400, content={"error": "Please upload an Excel file (.xlsx) for faculty summary"})
        
        # Save uploaded summary file temporarily
        temp_summary_path = "temp_summary.xlsx"
        summary_content = summary_file.file.read()
        with open(temp_summary_path, "wb") as f:
            f.write(summary_content)
        
        # Read the summary Excel file
        summary_df = pd.read_excel(temp_summary_path, engine="openpyxl")
        print(f"Summary file loaded with shape: {summary_df.shape}")
        print(f"Summary columns: {summary_df.columns.tolist()}")
        
        # Validate required columns for summary
        required_columns = ['Faculty', 'First Half Duties', 'Second Half Duties', 'First Half Dates', 'Second Half Dates']
        missing_columns = [col for col in required_columns if col not in summary_df.columns]
        if missing_columns:
            os.remove(temp_summary_path)
            return JSONResponse(status_code=400, content={"error": f"Missing required columns in faculty summary: {missing_columns}"})
        
        # Process exam schedule file if provided
        new_schedule = []
        if schedule_file and schedule_file.filename.endswith('.xlsx'):
            temp_schedule_path = "temp_schedule.xlsx"
            schedule_content = schedule_file.file.read()
            with open(temp_schedule_path, "wb") as f:
                f.write(schedule_content)
            
            try:
                schedule_df = pd.read_excel(temp_schedule_path, engine="openpyxl")
                
                # Validate schedule columns (expected: Date, First Half, Second Half)
                schedule_columns = schedule_df.columns.tolist()
                if 'Date' in schedule_columns:
                    for _, row in schedule_df.iterrows():
                        date_str = str(row['Date'])
                        if date_str and date_str != 'nan':
                            try:
                                # Try to parse the date
                                if '-' in date_str:
                                    date_obj = pd.to_datetime(date_str).date()
                                else:
                                    date_obj = datetime.datetime.strptime(date_str, '%d-%m-%Y').date()
                                
                                new_schedule.append({
                                    "date": date_obj.strftime('%Y-%m-%d'),
                                    "first_half": int(row.get('First Half', 0)) if pd.notna(row.get('First Half')) else 0,
                                    "second_half": int(row.get('Second Half', 0)) if pd.notna(row.get('Second Half')) else 0
                                })
                            except (ValueError, TypeError) as e:
                                print(f"Warning: Could not parse date '{date_str}' in schedule: {e}")
                
                # Save new schedule if valid data was found
                if new_schedule:
                    with open(EXAM_SCHEDULE_PATH, "w", encoding="utf-8") as f:
                        json.dump(new_schedule, f, ensure_ascii=False)
                    print(f"Updated exam schedule with {len(new_schedule)} dates")
                
                os.remove(temp_schedule_path)
            except Exception as e:
                print(f"Warning: Could not process schedule file: {e}")
                if os.path.exists(temp_schedule_path):
                    os.remove(temp_schedule_path)
        
        # Process faculty unavailability file if provided
        new_unavailability = {}
        if unavailability_file and unavailability_file.filename.endswith('.xlsx'):
            temp_unavailability_path = "temp_unavailability.xlsx"
            unavailability_content = unavailability_file.file.read()
            with open(temp_unavailability_path, "wb") as f:
                f.write(unavailability_content)
            
            try:
                unavailability_df = pd.read_excel(temp_unavailability_path, engine="openpyxl")
                
                # Expected format: Faculty, Date, Shift (First Half/Second Half)
                # or Faculty, First Half Dates, Second Half Dates
                unavailability_columns = unavailability_df.columns.tolist()
                
                if 'Faculty' in unavailability_columns:
                    for _, row in unavailability_df.iterrows():
                        faculty = str(row['Faculty'])
                        if faculty and faculty != 'nan':
                            if faculty not in new_unavailability:
                                new_unavailability[faculty] = {'first_half': [], 'second_half': []}
                            
                            # Handle different possible formats
                            if 'Date' in unavailability_columns and 'Shift' in unavailability_columns:
                                # Format: Faculty, Date, Shift
                                date_str = str(row['Date'])
                                shift = str(row['Shift'])
                                if date_str and date_str != 'nan' and shift and shift != 'nan':
                                    try:
                                        if '-' in date_str:
                                            date_obj = pd.to_datetime(date_str).date()
                                        else:
                                            date_obj = datetime.datetime.strptime(date_str, '%d-%m-%Y').date()
                                        
                                        date_formatted = date_obj.strftime('%Y-%m-%d')
                                        if 'First Half' in shift:
                                            if date_formatted not in new_unavailability[faculty]['first_half']:
                                                new_unavailability[faculty]['first_half'].append(date_formatted)
                                        elif 'Second Half' in shift:
                                            if date_formatted not in new_unavailability[faculty]['second_half']:
                                                new_unavailability[faculty]['second_half'].append(date_formatted)
                                    except (ValueError, TypeError) as e:
                                        print(f"Warning: Could not parse date '{date_str}' for faculty {faculty}: {e}")
                            
                            elif 'First Half Dates' in unavailability_columns and 'Second Half Dates' in unavailability_columns:
                                # Format: Faculty, First Half Dates, Second Half Dates
                                fh_dates_str = str(row['First Half Dates'])
                                sh_dates_str = str(row['Second Half Dates'])
                                
                                # Process First Half dates
                                if fh_dates_str and fh_dates_str != 'nan' and fh_dates_str != '':
                                    fh_dates = [d.strip() for d in fh_dates_str.split(',') if d.strip()]
                                    for date_str in fh_dates:
                                        try:
                                            date_obj = datetime.datetime.strptime(date_str, '%d-%m-%Y').date()
                                            date_formatted = date_obj.strftime('%Y-%m-%d')
                                            if date_formatted not in new_unavailability[faculty]['first_half']:
                                                new_unavailability[faculty]['first_half'].append(date_formatted)
                                        except ValueError:
                                            print(f"Warning: Could not parse first half date '{date_str}' for faculty {faculty}")
                                
                                # Process Second Half dates
                                if sh_dates_str and sh_dates_str != 'nan' and sh_dates_str != '':
                                    sh_dates = [d.strip() for d in sh_dates_str.split(',') if d.strip()]
                                    for date_str in sh_dates:
                                        try:
                                            date_obj = datetime.datetime.strptime(date_str, '%d-%m-%Y').date()
                                            date_formatted = date_obj.strftime('%Y-%m-%d')
                                            if date_formatted not in new_unavailability[faculty]['second_half']:
                                                new_unavailability[faculty]['second_half'].append(date_formatted)
                                        except ValueError:
                                            print(f"Warning: Could not parse second half date '{date_str}' for faculty {faculty}")
                
                # Save new unavailability if valid data was found
                if new_unavailability:
                    with open(FACULTY_UNAVAILABILITY_PATH, "w", encoding="utf-8") as f:
                        json.dump(new_unavailability, f, ensure_ascii=False)
                    print(f"Updated faculty unavailability for {len(new_unavailability)} faculty")
                
                os.remove(temp_unavailability_path)
            except Exception as e:
                print(f"Warning: Could not process unavailability file: {e}")
                if os.path.exists(temp_unavailability_path):
                    os.remove(temp_unavailability_path)
        
        # Generate new assignments from summary
        new_assignments = []
        slot_counter = {}  # (date, shift) -> list of faculty
        all_dates_shifts = set()
        for _, row in summary_df.iterrows():
            faculty = row['Faculty']
            # Process First Half dates
            fh_dates_str = str(row['First Half Dates'])
            if fh_dates_str and fh_dates_str != 'nan' and fh_dates_str != '':
                fh_dates = [d.strip() for d in fh_dates_str.split(',') if d.strip()]
                for date_str in fh_dates:
                    try:
                        date_obj = datetime.datetime.strptime(date_str, '%d-%m-%Y').date()
                        date_key = date_obj.strftime('%Y-%m-%d')
                        new_assignments.append({
                            "date": date_key,
                            "shift": "First Half",
                            "faculty": faculty
                        })
                        all_dates_shifts.add((date_key, "First Half"))
                        slot_counter.setdefault((date_key, "First Half"), []).append(faculty)
                    except ValueError:
                        print(f"Warning: Could not parse date '{date_str}' for faculty {faculty}")
            # Process Second Half dates
            sh_dates_str = str(row['Second Half Dates'])
            if sh_dates_str and sh_dates_str != 'nan' and sh_dates_str != '':
                sh_dates = [d.strip() for d in sh_dates_str.split(',') if d.strip()]
                for date_str in sh_dates:
                    try:
                        date_obj = datetime.datetime.strptime(date_str, '%d-%m-%Y').date()
                        date_key = date_obj.strftime('%Y-%m-%d')
                        new_assignments.append({
                            "date": date_key,
                            "shift": "Second Half",
                            "faculty": faculty
                        })
                        all_dates_shifts.add((date_key, "Second Half"))
                        slot_counter.setdefault((date_key, "Second Half"), []).append(faculty)
                    except ValueError:
                        print(f"Warning: Could not parse date '{date_str}' for faculty {faculty}")
        # Build schedule from summary
        new_schedule = []
        for (date, shift) in sorted(all_dates_shifts):
            required = len(slot_counter[(date, shift)])
            new_schedule.append({
                "date": date,
                "first_half": required if shift == "First Half" else 0,
                "second_half": required if shift == "Second Half" else 0
            })
        # Merge first_half and second_half for same date
        merged_schedule = {}
        for item in new_schedule:
            date = item["date"]
            if date not in merged_schedule:
                merged_schedule[date] = {"date": date, "first_half": 0, "second_half": 0}
            merged_schedule[date]["first_half"] += item["first_half"]
            merged_schedule[date]["second_half"] += item["second_half"]
        merged_schedule_list = list(merged_schedule.values())
        if merged_schedule_list:
            with open(EXAM_SCHEDULE_PATH, "w", encoding="utf-8") as f:
                json.dump(merged_schedule_list, f, ensure_ascii=False)
            print(f"Updated exam schedule from summary with {len(merged_schedule_list)} dates")
        # If no unavailability file, read from summary columns
        new_unavailability = {}
        if not (unavailability_file and unavailability_file.filename.endswith('.xlsx')):
            for _, row in summary_df.iterrows():
                faculty = row['Faculty']
                if faculty not in new_unavailability:
                    new_unavailability[faculty] = {'first_half': [], 'second_half': []}
                # First Half Unavailable
                fh_unavail = str(row.get('First Half Unavailable', ''))
                if fh_unavail and fh_unavail != 'None' and fh_unavail != 'nan':
                    fh_dates = [d.strip() for d in fh_unavail.split(',') if d.strip()]
                    for date_str in fh_dates:
                        try:
                            date_obj = datetime.datetime.strptime(date_str, '%d-%m-%Y').date()
                            date_key = date_obj.strftime('%Y-%m-%d')
                            if date_key not in new_unavailability[faculty]['first_half']:
                                new_unavailability[faculty]['first_half'].append(date_key)
                        except ValueError:
                            print(f"Warning: Could not parse unavailable first half date '{date_str}' for faculty {faculty}")
                # Second Half Unavailable
                sh_unavail = str(row.get('Second Half Unavailable', ''))
                if sh_unavail and sh_unavail != 'None' and sh_unavail != 'nan':
                    sh_dates = [d.strip() for d in sh_unavail.split(',') if d.strip()]
                    for date_str in sh_dates:
                        try:
                            date_obj = datetime.datetime.strptime(date_str, '%d-%m-%Y').date()
                            date_key = date_obj.strftime('%Y-%m-%d')
                            if date_key not in new_unavailability[faculty]['second_half']:
                                new_unavailability[faculty]['second_half'].append(date_key)
                        except ValueError:
                            print(f"Warning: Could not parse unavailable second half date '{date_str}' for faculty {faculty}")
            if new_unavailability:
                with open(FACULTY_UNAVAILABILITY_PATH, "w", encoding="utf-8") as f:
                    json.dump(new_unavailability, f, ensure_ascii=False)
                print(f"Updated faculty unavailability from summary for {len(new_unavailability)} faculty")
        
        # Save new assignments
        with open(ASSIGNMENTS_PATH, "w", encoding="utf-8") as f:
            json.dump(new_assignments, f, ensure_ascii=False)
        
        print(f"Generated {len(new_assignments)} assignments")
        
        # Clean up temp files
        os.remove(temp_summary_path)
        
        schedule_message = f" and updated exam schedule with {len(merged_schedule_list)} dates" if merged_schedule_list else ""
        unavailability_message = f" and updated unavailability for {len(new_unavailability)} faculty" if new_unavailability else ""
        
        print("Regeneration completed successfully")
        
        return {
            "status": "ok",
            "message": f"Regenerated {len(new_assignments)} assignments from summary{schedule_message}{unavailability_message}"
        }
        
    except Exception as e:
        print(f"Error regenerating from summary: {str(e)}")
        import traceback
        traceback.print_exc()
        
        # Clean up any temp files that might exist
        for temp_file in ["temp_summary.xlsx", "temp_schedule.xlsx", "temp_unavailability.xlsx"]:
            if os.path.exists(temp_file):
                try:
                    os.remove(temp_file)
                except:
                    pass
        
        return JSONResponse(status_code=500, content={"error": f"Failed to regenerate assignments: {str(e)}"})

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000) 